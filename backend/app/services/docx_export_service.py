from __future__ import annotations

import io
import uuid
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import Settings, settings
from app.core.logging import get_logger
from app.db.models.learning_output import LearningOutput
from app.db.models.meeting import Meeting
from app.db.models.meeting_insights import MeetingInsights
from app.db.models.question import Question
from app.db.models.transcript import Transcript
from app.services.storage_service import LocalTranscriptStorage
from app.db.repositories import learning_outputs as learning_output_repo
from app.db.repositories import questions as question_repo

logger = get_logger(__name__)


class DocxExportError(Exception):
    pass


class DocxExportService:
    def __init__(
        self,
        db: Session,
        *,
        config: Settings = settings,
        storage: LocalTranscriptStorage | None = None,
    ) -> None:
        self.db = db
        self.config = config
        self.storage = storage or LocalTranscriptStorage(config=config)

    def generate_docx(
        self,
        transcript_id: uuid.UUID,
        *,
        mcq_filters: dict | None = None,
        flashcard_filters: dict | None = None,
        short_question_filters: dict | None = None,
    ) -> Path:
        transcript = self.db.get(Transcript, transcript_id)
        if transcript is None:
            raise DocxExportError(f"Transcript not found: {transcript_id}")

        meeting = None
        if transcript.meeting_id is not None:
            meeting = self.db.get(Meeting, transcript.meeting_id)

        insights = self.db.scalar(
            select(MeetingInsights).where(MeetingInsights.transcript_id == transcript_id)
        )

        # Use existing filtering repositories to respect UI filters and ordering
        _flashcard_filters = flashcard_filters or {}
        _sq_filters = short_question_filters or {}
        _mcq_filters = mcq_filters or {}

        # Flashcards: reuse existing learning_output_repo with filter params
        fc_rows, _ = learning_output_repo.list_by_transcript(
            self.db,
            transcript_id,
            output_type="flashcard",
            category=_flashcard_filters.get("category"),
            difficulty=_flashcard_filters.get("difficulty"),
            offset=0,
            limit=_flashcard_filters.get("top", 10_000),
            order_desc=False,
            order_by_educational=bool(_flashcard_filters.get("top")),
        )
        flashcard_outputs = list(fc_rows)

        # Short Questions: reuse existing learning_output_repo with filter params
        sq_rows, _ = learning_output_repo.list_by_transcript(
            self.db,
            transcript_id,
            output_type="short_question",
            category=_sq_filters.get("category"),
            difficulty=_sq_filters.get("difficulty"),
            bloom_taxonomy=_sq_filters.get("bloom"),
            offset=0,
            limit=_sq_filters.get("top", 10_000),
            order_desc=False,
            order_by_educational=bool(_sq_filters.get("top")),
        )
        short_question_outputs = list(sq_rows)

        # MCQs: reuse existing question_repo with filter params
        mcq_difficulty = _mcq_filters.get("difficulty")
        mcq_category = _mcq_filters.get("category")
        mcq_bloom = _mcq_filters.get("bloom")
        mcq_top = _mcq_filters.get("top")
        mcq_rows, _ = question_repo.list_questions_for_transcript(
            self.db,
            transcript_id,
            question_type="mcq",
            difficulty=mcq_difficulty,
            category=mcq_category,
            bloom_taxonomy=mcq_bloom,
            offset=0,
            limit=mcq_top if mcq_top else 10_000,
            order_desc=False,
            order_by_educational=bool(mcq_top),
        )
        mcq_questions = list(mcq_rows)

        doc = Document()

        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        self._add_metadata_section(doc, transcript, meeting)
        self._add_insights_section(doc, insights)
        self._add_learning_outputs_section(doc, flashcard_outputs, short_question_outputs)
        self._add_questions_section(doc, mcq_questions)

        meeting_id_str = str(transcript.meeting_id) if transcript.meeting_id else "no-meeting"
        export_dir = self.storage.base_dir / "exports" / meeting_id_str
        export_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"transcript_{transcript_id.hex[:8]}_{timestamp}.docx"
        file_path = export_dir / filename

        doc.save(str(file_path))

        logger.info(
            "docx_export.generated",
            extra={
                "transcript_id": str(transcript_id),
                "file_path": str(file_path),
            },
        )

        return file_path

    def _add_metadata_section(
        self, doc: Document, transcript: Transcript, meeting: Meeting | None
    ) -> None:
        doc.add_heading("Transcript Export", level=0)

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Shading Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"

        meta_rows = [
            ("Transcript ID", str(transcript.id)),
            ("Meeting Topic", meeting.topic if meeting and meeting.topic else "N/A"),
            ("Meeting ID", str(transcript.meeting_id) if transcript.meeting_id else "N/A"),
            ("Source Format", transcript.source_format or "N/A"),
            ("Status", transcript.status),
            ("Language", transcript.language or "N/A"),
            ("Segment Count", str(transcript.segment_count) if transcript.segment_count else "N/A"),
            ("Word Count", str(transcript.word_count) if transcript.word_count else "N/A"),
            ("Generation Model", transcript.generation_model or "N/A"),
            ("Question Count", str(transcript.question_count) if transcript.question_count else "0"),
            ("Created At", str(transcript.created_at) if transcript.created_at else "N/A"),
        ]
        if meeting:
            meta_rows.append(("Duration (minutes)", str(meeting.duration_minutes) if meeting.duration_minutes else "N/A"))
            meta_rows.append(("Host", meeting.host_email or meeting.host_id or "N/A"))
            meta_rows.append(("Start Time", str(meeting.start_time) if meeting.start_time else "N/A"))

        for field_name, value in meta_rows:
            row = table.add_row().cells
            row[0].text = field_name
            row[1].text = value

        doc.add_paragraph()

    @staticmethod
    def _derive_key_concepts(insights: MeetingInsights) -> list[dict]:
        if insights.key_concepts:
            return insights.key_concepts
        return [
            {"concept": t.get("topic", "") if isinstance(t, dict) else str(t),
             "description": t.get("relevance", "") if isinstance(t, dict) else "",
             "importance_order": i + 1}
            for i, t in enumerate(insights.topics or [])
        ]

    @staticmethod
    def _derive_action_items(insights: MeetingInsights) -> list[dict]:
        if insights.action_items:
            return insights.action_items
        items: list[dict] = []
        for d in insights.decisions or []:
            decision = d.get("decision", "") if isinstance(d, dict) else str(d)
            decided_by = d.get("decided_by", "") if isinstance(d, dict) else ""
            items.append({"item_text": decision, "assignee": decided_by or None, "priority": None, "due_date": None})
        for r in insights.recommendations or []:
            rec = r.get("recommendation", "") if isinstance(r, dict) else str(r)
            audience = r.get("target_audience", "") if isinstance(r, dict) else ""
            priority = r.get("priority", "") if isinstance(r, dict) else ""
            items.append({"item_text": rec, "assignee": audience or None, "priority": priority or None, "due_date": None})
        return items

    def _add_insights_section(self, doc: Document, insights: MeetingInsights | None) -> None:
        doc.add_heading("Meeting Insights", level=1)

        if insights is None:
            doc.add_paragraph("No insights available for this transcript.")
            return

        derived_kc = self._derive_key_concepts(insights)
        derived_ai = self._derive_action_items(insights)

        if insights.summary_text:
            doc.add_heading("Summary", level=2)
            doc.add_paragraph(insights.summary_text)

        if derived_kc:
            doc.add_heading("Key Concepts", level=2)
            for kc in derived_kc:
                concept = kc.get("concept", "") if isinstance(kc, dict) else str(kc)
                desc = kc.get("description", "") if isinstance(kc, dict) else ""
                text = f"{concept}"
                if desc:
                    text += f" — {desc}"
                doc.add_paragraph(text, style="List Bullet")

        if derived_ai:
            doc.add_heading("Action Items", level=2)
            for ai in derived_ai:
                item = ai.get("item_text", "") if isinstance(ai, dict) else str(ai)
                assignee = ai.get("assignee", "") if isinstance(ai, dict) else ""
                priority = ai.get("priority", "") if isinstance(ai, dict) else ""
                text = item
                if assignee:
                    text += f" (Assignee: {assignee})"
                if priority:
                    text += f" [Priority: {priority}]"
                doc.add_paragraph(text, style="List Bullet")

        if insights.key_takeaways:
            doc.add_heading("Key Takeaways", level=2)
            for kt in insights.key_takeaways:
                takeaway = kt.get("takeaway", "") if isinstance(kt, dict) else str(kt)
                context = kt.get("context", "") if isinstance(kt, dict) else ""
                text = takeaway
                if context:
                    text += f" — {context}"
                doc.add_paragraph(text, style="List Bullet")

        if insights.learning_outcomes:
            doc.add_heading("Learning Outcomes", level=2)
            for lo in insights.learning_outcomes:
                outcome = lo.get("outcome", "") if isinstance(lo, dict) else str(lo)
                category = lo.get("category", "") if isinstance(lo, dict) else ""
                text = outcome
                if category:
                    text += f" [{category}]"
                doc.add_paragraph(text, style="List Bullet")

        if insights.topics:
            doc.add_heading("Topics", level=2)
            for t in insights.topics:
                topic = t.get("topic", "") if isinstance(t, dict) else str(t)
                relevance = t.get("relevance", "") if isinstance(t, dict) else ""
                text = topic
                if relevance:
                    text += f" (Relevance: {relevance})"
                doc.add_paragraph(text, style="List Bullet")

        if insights.decisions:
            doc.add_heading("Decisions", level=2)
            for d in insights.decisions:
                decision = d.get("decision", "") if isinstance(d, dict) else str(d)
                rationale = d.get("rationale", "") if isinstance(d, dict) else ""
                decided_by = d.get("decided_by", "") if isinstance(d, dict) else ""
                text = decision
                if decided_by:
                    text += f" (by {decided_by})"
                if rationale:
                    text += f" — {rationale}"
                doc.add_paragraph(text, style="List Bullet")

        if insights.recommendations:
            doc.add_heading("Recommendations", level=2)
            for r in insights.recommendations:
                rec = r.get("recommendation", "") if isinstance(r, dict) else str(r)
                priority = r.get("priority", "") if isinstance(r, dict) else ""
                audience = r.get("target_audience", "") if isinstance(r, dict) else ""
                text = rec
                if priority:
                    text += f" [Priority: {priority}]"
                if audience:
                    text += f" (For: {audience})"
                doc.add_paragraph(text, style="List Bullet")

        doc.add_paragraph()

    def _add_learning_outputs_section(
        self,
        doc: Document,
        flashcards: list[LearningOutput],
        short_questions: list[LearningOutput],
    ) -> None:
        doc.add_heading("Learning Outputs", level=1)

        if not flashcards and not short_questions:
            doc.add_paragraph("No learning outputs available for this transcript.")
            return

        if flashcards:
            doc.add_heading("Flashcards", level=2)
            for i, fc in enumerate(flashcards, 1):
                content = fc.content if isinstance(fc.content, dict) else {}
                front = content.get("front", "N/A")
                back = content.get("back", "N/A")
                category = content.get("category", "")

                p = doc.add_paragraph()
                run = p.add_run(f"#{i}: {front}")
                run.bold = True
                p.add_run(f"\nAnswer: {back}")
                if category:
                    p.add_run(f"\nCategory: {category}")

        if short_questions:
            doc.add_heading("Short Questions", level=2)
            for i, sq in enumerate(short_questions, 1):
                content = sq.content if isinstance(sq.content, dict) else {}
                question_text = content.get("question_text", "N/A")
                sample_answer = content.get("sample_answer", "")
                difficulty = sq.difficulty or ""

                p = doc.add_paragraph()
                run = p.add_run(f"#{i}: {question_text}")
                run.bold = True
                if sample_answer:
                    p.add_run(f"\nSample Answer: {sample_answer}")
                if difficulty:
                    p.add_run(f"\nDifficulty: {difficulty}")

        doc.add_paragraph()

    def _add_questions_section(self, doc: Document, mcqs: list[Question]) -> None:
        doc.add_heading("Generated Questions", level=1)

        if not mcqs:
            doc.add_paragraph("No questions available for this transcript.")
            return

        if mcqs:
            doc.add_heading("Multiple Choice Questions", level=2)
            for i, q in enumerate(mcqs, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"Q{i}: {q.question_text}")
                run.bold = True

                for opt in q.options:
                    if isinstance(opt, dict):
                        label = opt.get("label", "")
                        text = opt.get("text", "")
                    else:
                        label = ""
                        text = str(opt)
                    if label:
                        doc.add_paragraph(f"{label}. {text}", style="List Bullet")
                    else:
                        doc.add_paragraph(text, style="List Bullet")

                answer_p = doc.add_paragraph()
                answer_run = answer_p.add_run(f"Correct Answer: {q.correct_answer}")
                answer_run.bold = True
                answer_p.add_run(f"\nExplanation: {q.explanation}")
                if q.difficulty:
                    answer_p.add_run(f"\nDifficulty: {q.difficulty}")

        doc.add_paragraph()
